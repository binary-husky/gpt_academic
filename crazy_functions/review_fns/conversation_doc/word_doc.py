import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime
import docx
from docx.oxml import shared
from crazy_functions.doc_fns.conversation_doc.word_doc import convert_markdown_to_word


class WordFormatter:
    """聊天记录Word文档生成器 - 符合中国政府公文格式规范(GB/T 9704-2012)"""

    def __init__(self):
        self.doc = Document()
        self._setup_document()
        self._create_styles()

    def _setup_document(self):
        """设置文档基本格式，包括页面设置和页眉"""
        sections = self.doc.sections
        for section in sections:
            # 设置页面大小为A4
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # 设置页边距
            section.top_margin = Cm(3.7)  # 上边距37mm
            section.bottom_margin = Cm(3.5)  # 下边距35mm
            section.left_margin = Cm(2.8)  # 左边距28mm
            section.right_margin = Cm(2.6)  # 右边距26mm
            # 设置页眉页脚距离
            section.header_distance = Cm(2.0)
            section.footer_distance = Cm(2.0)

            # 修改页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = header_para.add_run("GPT-Academic学术对话 (体验地址：https://auth.gpt-academic.top/)")
            header_run.font.name = '仿宋'
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            header_run.font.size = Pt(9)

    def _create_styles(self):
        """创建文档样式"""
        # 创建正文样式
        style = self.doc.styles.add_style('Normal_Custom', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)

        # 创建问题样式
        question_style = self.doc.styles.add_style('Question_Style', WD_STYLE_TYPE.PARAGRAPH)
        question_style.font.name = '黑体'
        question_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        question_style.font.size = Pt(14)  # 调整为14磅
        question_style.font.bold = True
        question_style.paragraph_format.space_before = Pt(12)  # 减小段前距
        question_style.paragraph_format.space_after = Pt(6)
        question_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        question_style.paragraph_format.left_indent = Pt(0)  # 移除左缩进

        # 创建回答样式
        answer_style = self.doc.styles.add_style('Answer_Style', WD_STYLE_TYPE.PARAGRAPH)
        answer_style.font.name = '仿宋'
        answer_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        answer_style.font.size = Pt(12)  # 调整为12磅
        answer_style.paragraph_format.space_before = Pt(6)
        answer_style.paragraph_format.space_after = Pt(12)
        answer_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        answer_style.paragraph_format.left_indent = Pt(0)  # 移除左缩进

        # 创建标题样式
        title_style = self.doc.styles.add_style('Title_Custom', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = '黑体'  # 改用黑体
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_style.font.size = Pt(22)  # 调整为22磅
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(24)
        title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 添加参考文献样式
        ref_style = self.doc.styles.add_style('Reference_Style', WD_STYLE_TYPE.PARAGRAPH)
        ref_style.font.name = '宋体'
        ref_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        ref_style.font.size = Pt(10.5)  # 参考文献使用小号字体
        ref_style.paragraph_format.space_before = Pt(3)
        ref_style.paragraph_format.space_after = Pt(3)
        ref_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        ref_style.paragraph_format.left_indent = Pt(21)
        ref_style.paragraph_format.first_line_indent = Pt(-21)

        # 添加参考文献标题样式
        ref_title_style = self.doc.styles.add_style('Reference_Title_Style', WD_STYLE_TYPE.PARAGRAPH)
        ref_title_style.font.name = '黑体'
        ref_title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        ref_title_style.font.size = Pt(16)  # 参考文献标题与问题同样大小
        ref_title_style.font.bold = True
        ref_title_style.paragraph_format.space_before = Pt(24)  # 增加段前距
        ref_title_style.paragraph_format.space_after = Pt(12)
        ref_title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def create_document(self, question: str, answer: str, ranked_papers: list = None):
        """写入聊天历史
        Args:
            question: str, 用户问题
            answer: str, AI回答
            ranked_papers: list, 排序后的论文列表
        """
        try:
            # 添加标题
            title_para = self.doc.add_paragraph(style='Title_Custom')
            title_run = title_para.add_run('GPT-Academic 对话记录')

            # 添加日期
            try:
                date_para = self.doc.add_paragraph()
                date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                date_run = date_para.add_run(datetime.now().strftime('%Y年%m月%d日'))
                date_run.font.name = '仿宋'
                date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                date_run.font.size = Pt(16)
            except Exception as e:
                print(f"添加日期失败: {str(e)}")
                raise

            self.doc.add_paragraph()  # 添加空行

            # 添加问答对话
            try:
                q_para = self.doc.add_paragraph(style='Question_Style')
                q_para.add_run('问题：').bold = True
                q_para.add_run(str(question))

                a_para = self.doc.add_paragraph(style='Answer_Style')
                a_para.add_run('回答：').bold = True
                a_para.add_run(convert_markdown_to_word(str(answer)))
            except Exception as e:
                print(f"添加问答对话失败: {str(e)}")
                raise

            # 添加参考文献部分
            if ranked_papers:
                try:
                    ref_title = self.doc.add_paragraph(style='Reference_Title_Style')
                    ref_title.add_run("参考文献")
                    
                    for idx, paper in enumerate(ranked_papers, 1):
                        try:
                            ref_para = self.doc.add_paragraph(style='Reference_Style')
                            ref_para.add_run(f'[{idx}] ').bold = True
                            
                            # 添加作者
                            authors = ', '.join(paper.authors[:3])
                            if len(paper.authors) > 3:
                                authors += ' et al.'
                            ref_para.add_run(f'{authors}. ')
                            
                            # 添加标题
                            title_run = ref_para.add_run(paper.title)
                            title_run.italic = True
                            if hasattr(paper, 'url') and paper.url:
                                try:
                                    title_run._element.rPr.rStyle = self._create_hyperlink_style()
                                    self._add_hyperlink(ref_para, paper.title, paper.url)
                                except Exception as e:
                                    print(f"添加超链接失败: {str(e)}")
                            
                            # 添加期刊/会议信息
                            if paper.venue_name:
                                ref_para.add_run(f'. {paper.venue_name}')
                            
                            # 添加年份
                            if paper.year:
                                ref_para.add_run(f', {paper.year}')
                            
                            # 添加DOI
                            if paper.doi:
                                ref_para.add_run('. ')
                                if "arxiv" in paper.url:
                                    doi_url = paper.doi
                                else:   
                                    doi_url = f'https://doi.org/{paper.doi}'
                                self._add_hyperlink(ref_para, f'DOI: {paper.doi}', doi_url)
                            
                            ref_para.add_run('.')
                        except Exception as e:
                            print(f"添加第 {idx} 篇参考文献失败: {str(e)}")
                            continue
                except Exception as e:
                    print(f"添加参考文献部分失败: {str(e)}")
                    raise

            return self.doc
        
        except Exception as e:
            print(f"Word文档创建失败: {str(e)}")
            import traceback
            print(f"详细错误信息: {traceback.format_exc()}")
            raise

    def _create_hyperlink_style(self):
        """创建超链接样式"""
        styles = self.doc.styles
        if 'Hyperlink' not in styles:
            hyperlink_style = styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
            # 使用科技蓝 (#0066CC)
            hyperlink_style.font.color.rgb = 0x0066CC  # 科技蓝
            hyperlink_style.font.underline = True
        return styles['Hyperlink']

    def _add_hyperlink(self, paragraph, text, url):
        """添加超链接到段落"""
        # 这个是在XML级别添加超链接
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        
        # 创建超链接XML元素
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        
        # 创建文本运行
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        
        # 应用超链接样式
        rStyle = docx.oxml.shared.OxmlElement('w:rStyle')
        rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        
        # 添加文本
        t = docx.oxml.shared.OxmlElement('w:t')
        t.text = text
        new_run.append(rPr)
        new_run.append(t)
        hyperlink.append(new_run)
        
        # 将超链接添加到段落
        paragraph._p.append(hyperlink)

