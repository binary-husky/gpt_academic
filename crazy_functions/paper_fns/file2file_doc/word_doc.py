import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime

def convert_markdown_to_word(markdown_text):
    # 0. 首先标准化所有换行符为\n
    markdown_text = markdown_text.replace('\r\n', '\n').replace('\r', '\n')
    
    # 1. 处理标题 - 支持更多级别的标题，使用更精确的正则
    # 保留标题标记，以便后续处理时还能识别出标题级别
    markdown_text = re.sub(r'^(#{1,6})\s+(.+?)(?:\s+#+)?$', r'\1 \2', markdown_text, flags=re.MULTILINE)
    
    # 2. 处理粗体、斜体和加粗斜体
    markdown_text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', markdown_text)  # 加粗斜体
    markdown_text = re.sub(r'\*\*(.+?)\*\*', r'\1', markdown_text)      # 加粗
    markdown_text = re.sub(r'\*(.+?)\*', r'\1', markdown_text)          # 斜体
    markdown_text = re.sub(r'_(.+?)_', r'\1', markdown_text)            # 下划线斜体
    markdown_text = re.sub(r'__(.+?)__', r'\1', markdown_text)          # 下划线加粗
    
    # 3. 处理代码块 - 不移除，而是简化格式
    # 多行代码块
    markdown_text = re.sub(r'```(?:\w+)?\n([\s\S]*?)```', r'[代码块]\n\1[/代码块]', markdown_text)
    # 单行代码
    markdown_text = re.sub(r'`([^`]+)`', r'[代码]\1[/代码]', markdown_text)
    
    # 4. 处理列表 - 保留列表结构
    # 匹配无序列表
    markdown_text = re.sub(r'^(\s*)[-*+]\s+(.+?)$', r'\1• \2', markdown_text, flags=re.MULTILINE)
    
    # 5. 处理Markdown链接
    markdown_text = re.sub(r'\[([^\]]+)\]\(([^)]+?)\s*(?:"[^"]*")?\)', r'\1 (\2)', markdown_text)
    
    # 6. 处理HTML链接
    markdown_text = re.sub(r'<a href=[\'"]([^\'"]+)[\'"](?:\s+target=[\'"][^\'"]+[\'"])?>([^<]+)</a>', r'\2 (\1)', markdown_text)
    
    # 7. 处理图片
    markdown_text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[图片：\1]', markdown_text)
    
    return markdown_text


class WordFormatter:
    """文档Word格式化器 - 保留原始文档结构"""

    def __init__(self):
        self.doc = Document()
        self._setup_document()
        self._create_styles()

    def _setup_document(self):
        """设置文档基本格式，包括页面设置和页眉"""
        sections = self.doc.sections
        for section in sections:
            # 设置页面大小为A4
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # 设置页边距
            section.top_margin = Cm(3.7)  # 上边距37mm
            section.bottom_margin = Cm(3.5)  # 下边距35mm
            section.left_margin = Cm(2.8)  # 左边距28mm
            section.right_margin = Cm(2.6)  # 右边距26mm
            # 设置页眉页脚距离
            section.header_distance = Cm(2.0)
            section.footer_distance = Cm(2.0)

            # 添加页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header_run = header_para.add_run("文档处理结果")
            header_run.font.name = '仿宋'
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            header_run.font.size = Pt(9)

    def _create_styles(self):
        """创建文档样式"""
        # 创建正文样式
        style = self.doc.styles.add_style('Normal_Custom', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(12)  # 调整为12磅
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)

        # 创建标题样式
        title_style = self.doc.styles.add_style('Title_Custom', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = '黑体'
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_style.font.size = Pt(22)  # 调整为22磅
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(24)
        title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 创建标题1样式
        h1_style = self.doc.styles.add_style('Heading1_Custom', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = '黑体'
        h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)

        # 创建标题2样式
        h2_style = self.doc.styles.add_style('Heading2_Custom', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = '黑体'
        h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(6)

        # 创建标题3样式
        h3_style = self.doc.styles.add_style('Heading3_Custom', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = '黑体'
        h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)

        # 创建代码块样式
        code_style = self.doc.styles.add_style('Code_Custom', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(11)
        code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Pt(36)
        code_style.paragraph_format.right_indent = Pt(36)

        # 创建列表样式
        list_style = self.doc.styles.add_style('List_Custom', WD_STYLE_TYPE.PARAGRAPH)
        list_style.font.name = '仿宋'
        list_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        list_style.font.size = Pt(12)
        list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        list_style.paragraph_format.left_indent = Pt(21)
        list_style.paragraph_format.first_line_indent = Pt(-21)

    def create_document(self, content: str, processing_type: str = "文本处理"):
        """创建文档，保留原始结构"""
        # 添加标题
        title_para = self.doc.add_paragraph(style='Title_Custom')
        title_run = title_para.add_run('文档处理结果')

        # 添加处理类型
        processing_para = self.doc.add_paragraph()
        processing_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        processing_run = processing_para.add_run(f"处理方式: {processing_type}")
        processing_run.font.name = '仿宋'
        processing_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        processing_run.font.size = Pt(14)

        # 添加日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(f"处理时间: {datetime.now().strftime('%Y年%m月%d日')}")
        date_run.font.name = '仿宋'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        date_run.font.size = Pt(14)

        self.doc.add_paragraph()  # 添加空行

        # 预处理内容，将Markdown格式转换为适合Word的格式
        processed_content = convert_markdown_to_word(content)
        
        # 按行处理文本，保留结构
        lines = processed_content.split('\n')
        in_code_block = False
        current_paragraph = None
        
        for line in lines:
            # 检查是否为标题
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if header_match:
                # 根据#的数量确定标题级别
                level = len(header_match.group(1))
                title_text = header_match.group(2)
                
                if level == 1:
                    style = 'Heading1_Custom'
                elif level == 2:
                    style = 'Heading2_Custom'
                else:
                    style = 'Heading3_Custom'
                
                self.doc.add_paragraph(title_text, style=style)
                current_paragraph = None
            
            # 检查代码块标记
            elif '[代码块]' in line:
                in_code_block = True
                current_paragraph = self.doc.add_paragraph(style='Code_Custom')
                code_line = line.replace('[代码块]', '').strip()
                if code_line:
                    current_paragraph.add_run(code_line)
            
            elif '[/代码块]' in line:
                in_code_block = False
                code_line = line.replace('[/代码块]', '').strip()
                if code_line and current_paragraph:
                    current_paragraph.add_run(code_line)
                current_paragraph = None
            
            # 检查列表项
            elif line.strip().startswith('•'):
                p = self.doc.add_paragraph(style='List_Custom')
                p.add_run(line.strip())
                current_paragraph = None
            
            # 处理普通文本行
            elif line.strip():
                if in_code_block:
                    if current_paragraph:
                        current_paragraph.add_run('\n' + line)
                    else:
                        current_paragraph = self.doc.add_paragraph(line, style='Code_Custom')
                else:
                    if current_paragraph is None or not current_paragraph.text:
                        current_paragraph = self.doc.add_paragraph(line, style='Normal_Custom')
                    else:
                        current_paragraph.add_run('\n' + line)
            
            # 处理空行，创建新段落
            elif not in_code_block:
                current_paragraph = None

        return self.doc

