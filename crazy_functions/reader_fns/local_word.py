# encoding: utf-8
# @Time   : 2024/1/15
# @Author : Spike
# @Descr   : Word 文件处理工具

import os
import copy
# docx
from docx import Document
import xml.etree.ElementTree as ET


class DocxHandler:

    def __init__(self, docx_path, output_dir=None):
        """
        Args:
            docx_path: 读取文件路径
            output_dir: 保存路径，如果没有另存为操作，可以为空
        """
        if output_dir:
            self.output_dir = os.path.join(output_dir, 'word')
            os.makedirs(self.output_dir, exist_ok=True)
        self.markdown_content = ''
        try:
            self.doc = Document(docx_path)
        except Exception as e:
            raise ValueError(f"Error reading document: {e}")
        self.file_name = os.path.basename(docx_path).split('.')[0]

    def __extract_attribute_from_xml(self, xml_data, e_tag='blip', e_attr='embed') -> str:
        """从XML数据中提取属性"""
        root = ET.fromstring(xml_data)
        # 遍历所有元素
        for elem in root.iter():
            # 处理元素标签以忽略命名空间
            tag = elem.tag.split('}')[-1]  # 分割以去除命名空间
            if e_tag == tag:
                # 以相同的方式处理属性名称
                for attr in elem.attrib:
                    if e_attr in attr:
                        return elem.attrib[attr]

    def __is_inside_table(self, element):
        """检查元素是否位于表格内"""
        while element is not None:
            if element.tag.endswith('tbl'):
                return True
            element = element.getparent()
        return False

    def __get_markdown_heading_level(self, para_size):
        """根据字体大小确定Markdown标题的级别"""
        para_size = int(para_size)
        if para_size > 40:
            return 1
        elif para_size > 22:
            return 2
        else:
            return 0

    def __extract_and_save_image(self, inline, doc):
        """提取并保存内嵌图片"""
        image_path = ""
        # 检查inline对象是否含有图片
        if hasattr(inline, "graphic"):
            # 对于标准的内嵌图片
            graphic = inline.graphic
            if graphic and hasattr(graphic.graphicData, "pic") and hasattr(graphic.graphicData.pic, "blipFill"):
                blip = graphic.graphicData.pic.blipFill.blip
                if blip and hasattr(blip, "embed"):
                    image_part = doc.part.related_parts[blip.embed]
                    image_path = os.path.join(self.output_dir, f'{self.file_name}-image-{blip.embed}.png')
                    with open(image_path, 'wb') as f:
                        f.write(image_part.blob)
        else:
            # 对话内嵌表格的图片
            embed = self.__extract_attribute_from_xml(inline.xml)
            image_part = doc.part.related_parts[embed]
            image_path = os.path.join(self.output_dir, f'{self.file_name}-image-{embed}.png')
            with open(image_path, 'wb') as f:
                f.write(image_part.blob)

        # 返回Markdown图片链接，如果没有找到图片则返回空字符串
        return f"![{os.path.basename(image_path)}]({image_path})\n" if image_path else ""

    def _process_paragraph(self, element, doc):
        """处理段落中的文本和图片"""
        markdown_paragraph = ''
        para_text = ''.join(node.text for node in element.iter() if node.tag.endswith('t') and node.text)
        if para_text:
            para_size = self.__extract_attribute_from_xml(element.xml, e_tag='sz', e_attr='val')
            heading_level = self.__get_markdown_heading_level(para_size)
            heading_level = '#' * heading_level + ' ' if heading_level > 0 else ""
            markdown_paragraph += heading_level + para_text
        # 处理段落中的图片
        for inline in element.iter():
            if inline.tag.endswith('drawing'):
                image_markdown = self.__extract_and_save_image(inline, doc)
                markdown_paragraph += image_markdown

        return markdown_paragraph + '\n\n'

    def __process_table(self, element, doc):
        """将docx中的表格转换为Markdown格式，并处理表格中的图片"""
        tbl_element = copy.deepcopy(element)  # 深拷贝element以避免修改原始文档
        temp_doc = Document()
        temp_doc._element.body.append(tbl_element)
        temp_table = temp_doc.tables[-1]  # 获取最后一个表格，即刚刚添加的表格
        markdown = ""
        header_index = 0
        for row in temp_table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ''
                # 检查单元格中的文本和图片
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            cell_text += run.text.replace('\n', '<br>')  # 替换换行符
                        for inline in run.element.iter():
                            if inline.tag.endswith('drawing'):
                                image_markdown = self.__extract_and_save_image(inline, doc).replace('\n', '<br>')
                                cell_text += image_markdown
                row_data.append(cell_text)
            # 将表格行转换为Markdown格式
            markdown += "| " + " | ".join(row_data) + " |\n"
            if not header_index:
                headers = temp_table.rows[0].cells
                markdown += "| " + " | ".join(['---'] * len(headers)) + " |\n"
            header_index += 1
        return markdown

    def get_markdown(self):
        """提取docx文件、保留段落转换为Markdown格式"""
        for element in self.doc.element.body.iter():
            # 处理段落
            if element.tag.endswith('p') and not self.__is_inside_table(element):
                markdown_paragraph = self._process_paragraph(element, self.doc)
                self.markdown_content += markdown_paragraph

            # 处理表格
            elif element.tag.endswith('tbl'):
                table_markdown = self.__process_table(element, self.doc)
                self.markdown_content += table_markdown

        return self.markdown_content

    def save_markdown(self):
        """保存Markdown文件"""
        if not self.markdown_content:
            self.markdown_content = self.get_markdown()
        # Save Markdown file
        markdown_file_path = os.path.join(self.output_dir, f'{self.file_name}.md')
        try:
            with open(markdown_file_path, 'w', encoding='utf-8') as f:
                f.write(self.markdown_content)
            return markdown_file_path
        except IOError as e:
            raise IOError(f"Error writing markdown file: {e}")


if __name__ == '__main__':
    pass