import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime


def convert_markdown_to_word(markdown_text):
    # 0. 首先标准化所有换行符为\n
    markdown_text = markdown_text.replace('\r\n', '\n').replace('\r', '\n')

    # 1. 处理标题 - 支持更多级别的标题，使用更精确的正则
    # 保留标题标记，以便后续处理时还能识别出标题级别
    markdown_text = re.sub(r'^(#{1,6})\s+(.+?)(?:\s+#+)?$', r'\1 \2', markdown_text, flags=re.MULTILINE)

    # 2. 处理粗体、斜体和加粗斜体
    markdown_text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', markdown_text)  # 加粗斜体
    markdown_text = re.sub(r'\*\*(.+?)\*\*', r'\1', markdown_text)  # 加粗
    markdown_text = re.sub(r'\*(.+?)\*', r'\1', markdown_text)  # 斜体
    markdown_text = re.sub(r'_(.+?)_', r'\1', markdown_text)  # 下划线斜体
    markdown_text = re.sub(r'__(.+?)__', r'\1', markdown_text)  # 下划线加粗

    # 3. 处理代码块 - 不移除，而是简化格式
    # 多行代码块
    markdown_text = re.sub(r'```(?:\w+)?\n([\s\S]*?)```', r'[代码块]\n\1[/代码块]', markdown_text)
    # 单行代码
    markdown_text = re.sub(r'`([^`]+)`', r'[代码]\1[/代码]', markdown_text)

    # 4. 处理列表 - 保留列表结构
    # 匹配无序列表
    markdown_text = re.sub(r'^(\s*)[-*+]\s+(.+?)$', r'\1• \2', markdown_text, flags=re.MULTILINE)

    # 5. 处理Markdown链接
    markdown_text = re.sub(r'\[([^\]]+)\]\(([^)]+?)\s*(?:"[^"]*")?\)', r'\1 (\2)', markdown_text)

    # 6. 处理HTML链接
    markdown_text = re.sub(r'<a href=[\'"]([^\'"]+)[\'"](?:\s+target=[\'"][^\'"]+[\'"])?>([^<]+)</a>', r'\2 (\1)',
                           markdown_text)

    # 7. 处理图片
    markdown_text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[图片：\1]', markdown_text)

    return markdown_text


class WordFormatter:
    """聊天记录Word文档生成器 - 符合中国政府公文格式规范(GB/T 9704-2012)"""

    def __init__(self):
        self.doc = Document()
        self._setup_document()
        self._create_styles()

    def _setup_document(self):
        """设置文档基本格式，包括页面设置和页眉"""
        sections = self.doc.sections
        for section in sections:
            # 设置页面大小为A4
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # 设置页边距
            section.top_margin = Cm(3.7)  # 上边距37mm
            section.bottom_margin = Cm(3.5)  # 下边距35mm
            section.left_margin = Cm(2.8)  # 左边距28mm
            section.right_margin = Cm(2.6)  # 右边距26mm
            # 设置页眉页脚距离
            section.header_distance = Cm(2.0)
            section.footer_distance = Cm(2.0)

            # 添加页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header_run = header_para.add_run("GPT-Academic对话记录")
            header_run.font.name = '仿宋'
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            header_run.font.size = Pt(9)

    def _create_styles(self):
        """创建文档样式"""
        # 创建正文样式
        style = self.doc.styles.add_style('Normal_Custom', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(12)  # 调整为12磅
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)

        # 创建问题样式
        question_style = self.doc.styles.add_style('Question_Style', WD_STYLE_TYPE.PARAGRAPH)
        question_style.font.name = '黑体'
        question_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        question_style.font.size = Pt(14)  # 调整为14磅
        question_style.font.bold = True
        question_style.paragraph_format.space_before = Pt(12)  # 减小段前距
        question_style.paragraph_format.space_after = Pt(6)
        question_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        question_style.paragraph_format.left_indent = Pt(0)  # 移除左缩进

        # 创建回答样式
        answer_style = self.doc.styles.add_style('Answer_Style', WD_STYLE_TYPE.PARAGRAPH)
        answer_style.font.name = '仿宋'
        answer_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        answer_style.font.size = Pt(12)  # 调整为12磅
        answer_style.paragraph_format.space_before = Pt(6)
        answer_style.paragraph_format.space_after = Pt(12)
        answer_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        answer_style.paragraph_format.left_indent = Pt(0)  # 移除左缩进

        # 创建标题样式
        title_style = self.doc.styles.add_style('Title_Custom', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = '黑体'  # 改用黑体
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_style.font.size = Pt(22)  # 调整为22磅
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(24)
        title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # 添加参考文献样式
        ref_style = self.doc.styles.add_style('Reference_Style', WD_STYLE_TYPE.PARAGRAPH)
        ref_style.font.name = '宋体'
        ref_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        ref_style.font.size = Pt(10.5)  # 参考文献使用小号字体
        ref_style.paragraph_format.space_before = Pt(3)
        ref_style.paragraph_format.space_after = Pt(3)
        ref_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        ref_style.paragraph_format.left_indent = Pt(21)
        ref_style.paragraph_format.first_line_indent = Pt(-21)

        # 添加参考文献标题样式
        ref_title_style = self.doc.styles.add_style('Reference_Title_Style', WD_STYLE_TYPE.PARAGRAPH)
        ref_title_style.font.name = '黑体'
        ref_title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        ref_title_style.font.size = Pt(16)
        ref_title_style.font.bold = True
        ref_title_style.paragraph_format.space_before = Pt(24)
        ref_title_style.paragraph_format.space_after = Pt(12)
        ref_title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def create_document(self,  history):
        """写入聊天历史"""
        # 添加标题
        title_para = self.doc.add_paragraph(style='Title_Custom')
        title_run = title_para.add_run('GPT-Academic 对话记录')

        # 添加日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_para.add_run(datetime.now().strftime('%Y年%m月%d日'))
        date_run.font.name = '仿宋'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        date_run.font.size = Pt(16)

        self.doc.add_paragraph()  # 添加空行

        # 添加对话内容
        for i in range(0, len(history), 2):
            question = history[i]
            answer = convert_markdown_to_word(history[i + 1])

            if question:
                q_para = self.doc.add_paragraph(style='Question_Style')
                q_para.add_run(f'问题 {i//2 + 1}：').bold = True
                q_para.add_run(str(question))

            if answer:
                a_para = self.doc.add_paragraph(style='Answer_Style')
                a_para.add_run(f'回答 {i//2 + 1}：').bold = True
                a_para.add_run(str(answer))


        return self.doc

