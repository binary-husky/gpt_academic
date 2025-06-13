import os
import time
from abc import ABC, abstractmethod
from datetime import datetime
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import  Inches, Cm
from docx.shared import Pt, RGBColor, Inches
from typing import Dict, List, Tuple
import markdown
from crazy_functions.doc_fns.conversation_doc.word_doc import convert_markdown_to_word



class DocumentFormatter(ABC):
    """文档格式化基类，定义文档格式化的基本接口"""

    def __init__(self, final_summary: str, file_summaries_map: Dict, failed_files: List[Tuple]):
        self.final_summary = final_summary
        self.file_summaries_map = file_summaries_map
        self.failed_files = failed_files

    @abstractmethod
    def format_failed_files(self) -> str:
        """格式化失败文件列表"""
        pass

    @abstractmethod
    def format_file_summaries(self) -> str:
        """格式化文件总结内容"""
        pass

    @abstractmethod
    def create_document(self) -> str:
        """创建完整文档"""
        pass


class WordFormatter(DocumentFormatter):
    """Word格式文档生成器 - 符合中国政府公文格式规范(GB/T 9704-2012)，并进行了优化"""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.doc = Document()
        self._setup_document()
        self._create_styles()
        # 初始化三级标题编号系统
        self.numbers = {
            1: 0,  # 一级标题编号
            2: 0,  # 二级标题编号
            3: 0  # 三级标题编号
        }

    def _setup_document(self):
        """设置文档基本格式，包括页面设置和页眉"""
        sections = self.doc.sections
        for section in sections:
            # 设置页面大小为A4
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # 设置页边距
            section.top_margin = Cm(3.7)  # 上边距37mm
            section.bottom_margin = Cm(3.5)  # 下边距35mm
            section.left_margin = Cm(2.8)  # 左边距28mm
            section.right_margin = Cm(2.6)  # 右边距26mm
            # 设置页眉页脚距离
            section.header_distance = Cm(2.0)
            section.footer_distance = Cm(2.0)

            # 添加页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header_run = header_para.add_run("该文档由GPT-academic生成")
            header_run.font.name = '仿宋'
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            header_run.font.size = Pt(9)

    def _create_styles(self):
        """创建文档样式"""
        # 创建正文样式
        style = self.doc.styles.add_style('Normal_Custom', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Pt(28)

        # 创建各级标题样式
        self._create_heading_style('Title_Custom', '方正小标宋简体', 32, WD_PARAGRAPH_ALIGNMENT.CENTER)
        self._create_heading_style('Heading1_Custom', '黑体', 22, WD_PARAGRAPH_ALIGNMENT.LEFT)
        self._create_heading_style('Heading2_Custom', '黑体', 18, WD_PARAGRAPH_ALIGNMENT.LEFT)
        self._create_heading_style('Heading3_Custom', '黑体', 16, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def _create_heading_style(self, style_name: str, font_name: str, font_size: int, alignment):
        """创建标题样式"""
        style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        return style

    def _get_heading_number(self, level: int) -> str:
        """
        生成标题编号

        Args:
            level: 标题级别 (0-3)

        Returns:
            str: 格式化的标题编号
        """
        if level == 0:  # 主标题不需要编号
            return ""

        self.numbers[level] += 1  # 增加当前级别的编号

        # 重置下级标题编号
        for i in range(level + 1, 4):
            self.numbers[i] = 0

        # 根据级别返回不同格式的编号
        if level == 1:
            return f"{self.numbers[1]}. "
        elif level == 2:
            return f"{self.numbers[1]}.{self.numbers[2]} "
        elif level == 3:
            return f"{self.numbers[1]}.{self.numbers[2]}.{self.numbers[3]} "
        return ""

    def _add_heading(self, text: str, level: int):
        """
        添加带编号的标题

        Args:
            text: 标题文本
            level: 标题级别 (0-3)
        """
        style_map = {
            0: 'Title_Custom',
            1: 'Heading1_Custom',
            2: 'Heading2_Custom',
            3: 'Heading3_Custom'
        }

        number = self._get_heading_number(level)
        paragraph = self.doc.add_paragraph(style=style_map[level])

        if number:
            number_run = paragraph.add_run(number)
            font_size = 22 if level == 1 else (18 if level == 2 else 16)
            self._get_run_style(number_run, '黑体', font_size, True)

        text_run = paragraph.add_run(text)
        font_size = 32 if level == 0 else (22 if level == 1 else (18 if level == 2 else 16))
        self._get_run_style(text_run, '黑体', font_size, True)

        # 主标题添加日期
        if level == 0:
            date_paragraph = self.doc.add_paragraph()
            date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            date_run = date_paragraph.add_run(datetime.now().strftime('%Y年%m月%d日'))
            self._get_run_style(date_run, '仿宋', 16, False)

        return paragraph

    def _get_run_style(self, run, font_name: str, font_size: int, bold: bool = False):
        """设置文本运行对象的样式"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold

    def format_failed_files(self) -> str:
        """格式化失败文件列表"""
        result = []
        if not self.failed_files:
            return "\n".join(result)

        result.append("处理失败文件:")
        for fp, reason in self.failed_files:
            result.append(f"• {os.path.basename(fp)}: {reason}")

        self._add_heading("处理失败文件", 1)
        for fp, reason in self.failed_files:
            self._add_content(f"• {os.path.basename(fp)}: {reason}", indent=False)
        self.doc.add_paragraph()

        return "\n".join(result)

    def _add_content(self, text: str, indent: bool = True):
        """添加正文内容，使用convert_markdown_to_word处理文本"""
        # 使用convert_markdown_to_word处理markdown文本
        processed_text = convert_markdown_to_word(text)
        paragraph = self.doc.add_paragraph(processed_text, style='Normal_Custom')
        if not indent:
            paragraph.paragraph_format.first_line_indent = Pt(0)
        return paragraph

    def format_file_summaries(self) -> str:
        """
        格式化文件总结内容，确保正确的标题层级并处理markdown文本
        """
        result = []
        # 首先对文件路径进行分组整理
        file_groups = {}
        for path in sorted(self.file_summaries_map.keys()):
            dir_path = os.path.dirname(path)
            if dir_path not in file_groups:
                file_groups[dir_path] = []
            file_groups[dir_path].append(path)

        # 处理没有目录的文件
        root_files = file_groups.get("", [])
        if root_files:
            for path in sorted(root_files):
                file_name = os.path.basename(path)
                result.append(f"\n📄 {file_name}")
                result.append(self.file_summaries_map[path])
                # 无目录的文件作为二级标题
                self._add_heading(f"📄 {file_name}", 2)
                # 使用convert_markdown_to_word处理文件内容
                self._add_content(convert_markdown_to_word(self.file_summaries_map[path]))
                self.doc.add_paragraph()

        # 处理有目录的文件
        for dir_path in sorted(file_groups.keys()):
            if dir_path == "":  # 跳过已处理的根目录文件
                continue

            # 添加目录作为二级标题
            result.append(f"\n📁 {dir_path}")
            self._add_heading(f"📁 {dir_path}", 2)

            # 该目录下的所有文件作为三级标题
            for path in sorted(file_groups[dir_path]):
                file_name = os.path.basename(path)
                result.append(f"\n📄 {file_name}")
                result.append(self.file_summaries_map[path])

                # 添加文件名作为三级标题
                self._add_heading(f"📄 {file_name}", 3)
                # 使用convert_markdown_to_word处理文件内容
                self._add_content(convert_markdown_to_word(self.file_summaries_map[path]))
                self.doc.add_paragraph()

        return "\n".join(result)


    def create_document(self):
        """创建完整Word文档并返回文档对象"""
        # 重置所有编号
        for level in self.numbers:
            self.numbers[level] = 0

        # 添加主标题
        self._add_heading("文档总结报告", 0)
        self.doc.add_paragraph()

        # 添加总体摘要，使用convert_markdown_to_word处理
        self._add_heading("总体摘要", 1)
        self._add_content(convert_markdown_to_word(self.final_summary))
        self.doc.add_paragraph()

        # 添加失败文件列表（如果有）
        if self.failed_files:
            self.format_failed_files()

        # 添加文件详细总结
        self._add_heading("各文件详细总结", 1)
        self.format_file_summaries()

        return self.doc

    def save_as_pdf(self, word_path, pdf_path=None):
        """将生成的Word文档转换为PDF
        
        参数:
            word_path: Word文档的路径
            pdf_path: 可选，PDF文件的输出路径。如果未指定，将使用与Word文档相同的名称和位置
            
        返回:
            生成的PDF文件路径，如果转换失败则返回None
        """
        from crazy_functions.doc_fns.conversation_doc.word2pdf import WordToPdfConverter
        try:
            pdf_path = WordToPdfConverter.convert_to_pdf(word_path, pdf_path)
            return pdf_path
        except Exception as e:
            print(f"PDF转换失败: {str(e)}")
            return None


class MarkdownFormatter(DocumentFormatter):
    """Markdown格式文档生成器"""

    def format_failed_files(self) -> str:
        if not self.failed_files:
            return ""

        formatted_text = ["\n## ⚠️ 处理失败的文件"]
        for fp, reason in self.failed_files:
            formatted_text.append(f"- {os.path.basename(fp)}: {reason}")
        formatted_text.append("\n---")
        return "\n".join(formatted_text)

    def format_file_summaries(self) -> str:
        formatted_text = []
        sorted_paths = sorted(self.file_summaries_map.keys())
        current_dir = ""

        for path in sorted_paths:
            dir_path = os.path.dirname(path)
            if dir_path != current_dir:
                if dir_path:
                    formatted_text.append(f"\n## 📁 {dir_path}")
                current_dir = dir_path

            file_name = os.path.basename(path)
            formatted_text.append(f"\n### 📄 {file_name}")
            formatted_text.append(self.file_summaries_map[path])
            formatted_text.append("\n---")

        return "\n".join(formatted_text)

    def create_document(self) -> str:
        document = [
            "# 📑 文档总结报告",
            "\n## 总体摘要",
            self.final_summary
        ]

        if self.failed_files:
            document.append(self.format_failed_files())

        document.extend([
            "\n# 📚 各文件详细总结",
            self.format_file_summaries()
        ])

        return "\n".join(document)



class HtmlFormatter(DocumentFormatter):
    """HTML格式文档生成器 - 优化版"""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.md = markdown.Markdown(extensions=['extra','codehilite', 'tables','nl2br'])
        self.css_styles = """
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }

        @keyframes slideIn {
            from { transform: translateX(-20px); opacity: 0; }
            to { transform: translateX(0); opacity: 1; }
        }

        @keyframes pulse {
            0% { transform: scale(1); }
            50% { transform: scale(1.05); }
            100% { transform: scale(1); }
        }

        :root {
            /* Enhanced color palette */
            --primary-color: #2563eb;
            --primary-light: #eff6ff;
            --secondary-color: #1e293b;
            --background-color: #f8fafc;
            --text-color: #334155;
            --text-light: #64748b;
            --border-color: #e2e8f0;
            --error-color: #ef4444;
            --error-light: #fef2f2;
            --success-color: #22c55e;
            --warning-color: #f59e0b;
            --card-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1), 0 2px 4px -2px rgb(0 0 0 / 0.1);
            --hover-shadow: 0 20px 25px -5px rgb(0 0 0 / 0.1), 0 8px 10px -6px rgb(0 0 0 / 0.1);
            
            /* Typography */
            --heading-font: "Plus Jakarta Sans", system-ui, sans-serif;
            --body-font: "Inter", system-ui, sans-serif;
        }

        body {
            font-family: var(--body-font);
            line-height: 1.8;
            max-width: 1200px;
            margin: 0 auto;
            padding: 2rem;
            color: var(--text-color);
            background-color: var(--background-color);
            font-size: 16px;
            -webkit-font-smoothing: antialiased;
        }

        .container {
            background: white;
            padding: 3rem;
            border-radius: 24px;
            box-shadow: var(--card-shadow);
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            animation: fadeIn 0.6s ease-out;
            border: 1px solid var(--border-color);
        }

        .container:hover {
            box-shadow: var(--hover-shadow);
            transform: translateY(-2px);
        }

        h1, h2, h3 {
            font-family: var(--heading-font);
            font-weight: 600;
        }

        h1 {
            color: var(--primary-color);
            font-size: 2.8em;
            text-align: center;
            margin: 2rem 0 3rem;
            padding-bottom: 1.5rem;
            border-bottom: 3px solid var(--primary-color);
            letter-spacing: -0.03em;
            position: relative;
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 1rem;
        }

        h1::after {
            content: '';
            position: absolute;
            bottom: -3px;
            left: 50%;
            transform: translateX(-50%);
            width: 120px;
            height: 3px;
            background: linear-gradient(90deg, var(--primary-color), var(--primary-light));
            border-radius: 3px;
            transition: width 0.3s ease;
        }

        h1:hover::after {
            width: 180px;
        }

        h2 {
            color: var(--secondary-color);
            font-size: 1.9em;
            margin: 2.5rem 0 1.5rem;
            padding-left: 1.2rem;
            border-left: 4px solid var(--primary-color);
            letter-spacing: -0.02em;
            display: flex;
            align-items: center;
            gap: 1rem;
            transition: all 0.3s ease;
        }

        h2:hover {
            color: var(--primary-color);
            transform: translateX(5px);
        }

        h3 {
            color: var(--text-color);
            font-size: 1.5em;
            margin: 2rem 0 1rem;
            padding-bottom: 0.8rem;
            border-bottom: 2px solid var(--border-color);
            transition: all 0.3s ease;
            display: flex;
            align-items: center;
            gap: 0.8rem;
        }

        h3:hover {
            color: var(--primary-color);
            border-bottom-color: var(--primary-color);
        }

        .summary {
            background: var(--primary-light);
            padding: 2.5rem;
            border-radius: 16px;
            margin: 2.5rem 0;
            box-shadow: 0 4px 6px -1px rgba(37, 99, 235, 0.1);
            position: relative;
            overflow: hidden;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
            animation: slideIn 0.5s ease-out;
        }

        .summary:hover {
            transform: translateY(-3px);
            box-shadow: 0 8px 12px -2px rgba(37, 99, 235, 0.15);
        }

        .summary::before {
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(to bottom, var(--primary-color), rgba(37, 99, 235, 0.6));
        }

        .summary p {
            margin: 1.2rem 0;
            line-height: 1.9;
            color: var(--text-color);
            transition: color 0.3s ease;
        }

        .summary:hover p {
            color: var(--secondary-color);
        }

        .details {
            margin-top: 3.5rem;
            padding-top: 2.5rem;
            border-top: 2px dashed var(--border-color);
            animation: fadeIn 0.8s ease-out;
        }

        .failed-files {
            background: var(--error-light);
            padding: 2rem;
            border-radius: 16px;
            margin: 3rem 0;
            border-left: 4px solid var(--error-color);
            position: relative;
            transition: all 0.3s ease;
            animation: slideIn 0.5s ease-out;
        }

        .failed-files:hover {
            transform: translateX(5px);
            box-shadow: 0 8px 15px -3px rgba(239, 68, 68, 0.1);
        }

        .failed-files h2 {
            color: var(--error-color);
            border-left: none;
            padding-left: 0;
        }

        .failed-files ul {
            margin: 1.8rem 0;
            padding-left: 1.2rem;
            list-style-type: none;
        }

        .failed-files li {
            margin: 1.2rem 0;
            padding: 1.2rem 1.8rem;
            background: rgba(239, 68, 68, 0.08);
            border-radius: 12px;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        }

        .failed-files li:hover {
            transform: translateX(8px);
            background: rgba(239, 68, 68, 0.12);
        }

        .directory-section {
            margin: 3.5rem 0;
            padding: 2rem;
            background: var(--background-color);
            border-radius: 16px;
            position: relative;
            transition: all 0.3s ease;
            animation: fadeIn 0.6s ease-out;
        }

        .directory-section:hover {
            background: white;
            box-shadow: var(--card-shadow);
        }

        .file-summary {
            background: white;
            padding: 2rem;
            margin: 1.8rem 0;
            border-radius: 16px;
            box-shadow: var(--card-shadow);
            border-left: 4px solid var(--border-color);
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }

        .file-summary:hover {
            border-left-color: var(--primary-color);
            transform: translateX(8px) translateY(-2px);
            box-shadow: var(--hover-shadow);
        }

        .file-summary {
            background: white;
            padding: 2rem;
            margin: 1.8rem 0;
            border-radius: 16px;
            box-shadow: var(--card-shadow);
            border-left: 4px solid var(--border-color);
            transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
        }

        .file-summary:hover {
            border-left-color: var(--primary-color);
            transform: translateX(8px) translateY(-2px);
            box-shadow: var(--hover-shadow);
        }

        .icon {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 32px;
            height: 32px;
            border-radius: 8px;
            background: var(--primary-light);
            color: var(--primary-color);
            font-size: 1.2em;
            transition: all 0.3s ease;
        }

        .file-summary:hover .icon,
        .directory-section:hover .icon {
            transform: scale(1.1);
            background: var(--primary-color);
            color: white;
        }

        /* Smooth scrolling */
        html {
            scroll-behavior: smooth;
        }

        /* Selection style */
        ::selection {
            background: var(--primary-light);
            color: var(--primary-color);
        }

        /* Print styles */
        @media print {
            body {
                background: white;
            }
            .container {
                box-shadow: none;
                padding: 0;
            }
            .file-summary, .failed-files {
                break-inside: avoid;
                box-shadow: none;
            }
            .icon {
                display: none;
            }
        }

        /* Responsive design */
        @media (max-width: 768px) {
            body {
                padding: 1rem;
                font-size: 15px;
            }
            
            .container {
                padding: 1.5rem;
            }

            h1 {
                font-size: 2.2em;
                margin: 1.5rem 0 2rem;
            }

            h2 {
                font-size: 1.7em;
            }

            h3 {
                font-size: 1.4em;
            }

            .summary, .failed-files, .directory-section {
                padding: 1.5rem;
            }

            .file-summary {
                padding: 1.2rem;
            }

            .icon {
                width: 28px;
                height: 28px;
            }
        }

        /* Dark mode support */
        @media (prefers-color-scheme: dark) {
            :root {
                --primary-light: rgba(37, 99, 235, 0.15);
                --background-color: #0f172a;
                --text-color: #e2e8f0;
                --text-light: #94a3b8;
                --border-color: #1e293b;
                --error-light: rgba(239, 68, 68, 0.15);
            }

            .container, .file-summary {
                background: #1e293b;
            }

            .directory-section {
                background: #0f172a;
            }

            .directory-section:hover {
                background: #1e293b;
            }
        }
        """

    def format_failed_files(self) -> str:
        if not self.failed_files:
            return ""

        failed_files_html = ['<div class="failed-files">']
        failed_files_html.append('<h2><span class="icon">⚠️</span> 处理失败的文件</h2>')
        failed_files_html.append("<ul>")
        for fp, reason in self.failed_files:
            failed_files_html.append(
                f'<li><strong>📄 {os.path.basename(fp)}</strong><br><span style="color: var(--text-light)">{reason}</span></li>'
            )
        failed_files_html.append("</ul></div>")
        return "\n".join(failed_files_html)

    def format_file_summaries(self) -> str:
        formatted_html = []
        sorted_paths = sorted(self.file_summaries_map.keys())
        current_dir = ""

        for path in sorted_paths:
            dir_path = os.path.dirname(path)
            if dir_path != current_dir:
                if dir_path:
                    formatted_html.append('<div class="directory-section">')
                    formatted_html.append(f'<h2><span class="icon">📁</span> {dir_path}</h2>')
                    formatted_html.append('</div>')
                current_dir = dir_path

            file_name = os.path.basename(path)
            formatted_html.append('<div class="file-summary">')
            formatted_html.append(f'<h3><span class="icon">📄</span> {file_name}</h3>')
            formatted_html.append(self.md.convert(self.file_summaries_map[path]))
            formatted_html.append('</div>')

        return "\n".join(formatted_html)

    def create_document(self) -> str:
        """生成HTML文档
        Returns:
            str: 完整的HTML文档字符串
        """
        return f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="utf-8">
            <meta name="viewport" content="width=device-width, initial-scale=1">
            <title>文档总结报告</title>
            <link href="https://cdnjs.cloudflare.com/ajax/libs/inter/3.19.3/inter.css" rel="stylesheet">
            <link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;600&display=swap" rel="stylesheet">
            <style>{self.css_styles}</style>
        </head>
        <body>
            <div class="container">
                <h1><span class="icon">📑</span> 文档总结报告</h1>
                <div class="summary">
                    <h2><span class="icon">📋</span> 总体摘要</h2>
                    <p>{self.md.convert(self.final_summary)}</p>
                </div>
                {self.format_failed_files()}
                <div class="details">
                    <h2><span class="icon">📚</span> 各文件详细总结</h2>
                    {self.format_file_summaries()}
                </div>
            </div>

        </body>
        </html>
        """